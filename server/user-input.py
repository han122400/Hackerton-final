# user_input.py (통합 테스트용)
import os
import json
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import JSONResponse
from supabase import create_client
from io import BytesIO
from docx import Document
from datetime import datetime

app = FastAPI()

# ======================= 🔹 Supabase 연결 ======================= #
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
TABLE = "interviews"

supabase = create_client(SUPABASE_URL, SUPABASE_KEY)

# ======================= 🔹 이력서 분석 함수 ======================= #
def extract_text_and_tables_from_docx(file_bytes):
    doc = Document(BytesIO(file_bytes))
    full_text, tables_data = [], []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    for table in doc.tables:
        tdata = []
        for row in table.rows:
            tdata.append([cell.text.strip() for cell in row.cells])
        tables_data.append(tdata)
    return full_text, tables_data

def extract_sections(full_text, tables_data):
    sections = {"학력":"", "가족사항":"", "경력":"", "프로젝트":"", "자격증":"", "기타":""}
    current_section = "기타"
    for line in full_text:
        line_clean = line.replace(" ","")
        if "학력" in line_clean: current_section="학력"
        elif "가족사항" in line_clean: current_section="가족사항"
        elif "경력" in line_clean: current_section="경력"
        elif "프로젝트" in line_clean: current_section="프로젝트"
        elif "자격증" in line_clean or "컴퓨터" in line_clean: current_section="자격증"
        sections[current_section] += line + "\n"
    for table in tables_data:
        header_row = table[0]
        header_text = "".join(header_row).replace(" ", "")
        target_section = "기타"
        if any(k in header_text for k in ["졸업","학교","학력"]): target_section="학력"
        elif any(k in header_text for k in ["관계","직업","가족"]): target_section="가족사항"
        elif any(k in header_text for k in ["근무","회사","직위","경력"]): target_section="경력"
        elif any(k in header_text for k in ["프로젝트"]): target_section="프로젝트"
        elif any(k in header_text for k in ["자격","컴퓨터"]): target_section="자격증"
        sections[target_section] += "\n".join(["\t".join(row) for row in table]) + "\n"
    return sections

# ======================= 🔹 API 엔드포인트 ======================= #
@app.post("/user-input")
async def save_user_input(
    position: str = Form(...),
    company: str = Form(None),
    notes: str = Form(None),
    userName: str = Form(None),
    resume: UploadFile = File(None),
):
    try:
        print("\n=== 새 요청 ===")
        print(f"position: {position}, company: {company}, userName: {userName}")

        # 1️⃣ 이력서 분석
        analysis_result = {}
        if resume:
            file_bytes = await resume.read()
            full_text, tables = extract_text_and_tables_from_docx(file_bytes)
            analysis_result = extract_sections(full_text, tables)
            print("✅ 분석 결과:")
            for k,v in analysis_result.items():
                preview = v.strip()[:200] + ("..." if len(v.strip()) > 200 else "")
                print(f"--- {k} ---\n{preview}\n")
        else:
            print("⚠️ 이력서 파일 없음")

        # 2️⃣ Supabase에 저장 (JSONB 직렬화 포함)
        record = {
            "user_name": userName,
            "position": position,
            "company": company,
            "notes": notes,
            "start_time": datetime.utcnow(),  # TIMESTAMP WITH TIME ZONE
            "analysis": json.dumps(analysis_result or {})  # JSONB 안전 저장
        }
        res = supabase.table(TABLE).insert(record).execute()

        if res.error:
            print("❌ DB 저장 실패:", res.error)
            return JSONResponse(status_code=400, content={"ok": False, "error": res.error.message})

        print("✅ DB 저장 성공")
        print("DB 응답:", res.data)

        return {"ok": True, "analysis": analysis_result}

    except Exception as e:
        print("❌ 예외 발생:", e)
        return JSONResponse(status_code=500, content={"ok": False, "error": str(e)})
